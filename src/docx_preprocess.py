from docx import Document
import re
import os
import json


from docx import Document


def generate_navigation_dict(document_path):
    document = Document(document_path)
    navigation_dict = {}
    current_heading = None
    current_text = []
    
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            if current_heading is not None:
                combined_text = ' '.join([current_heading[1].split('\t', 1)[-1]] + current_text)
                if len(combined_text.split()) >= 5:
                    navigation_dict[current_heading] = combined_text
                current_text = []
            
            # Check if the last character is a digit
            if paragraph.style.name[-1].isdigit():
                heading_level = int(paragraph.style.name[-1])
            else:
                # Use the preset heading level if the last character is not a digit
                heading_level = 1
            
            heading_text = paragraph.text
            current_heading = (heading_level, heading_text)
        else:
            if current_heading is not None and paragraph.text.strip():
                current_text.append(paragraph.text)
    
    if current_heading is not None:
        combined_text = ' '.join([current_heading[1].split('\t', 1)[-1]] + current_text)
        if len(combined_text.split()) >= 10:
            navigation_dict[current_heading] = combined_text
    
    return navigation_dict

def rel14_navigation_dict(document_path):
    document = Document(document_path)
    navigation_dict = {}
    current_heading = None
    current_text = []
    skipping = True

    skip_text = "The present document provides a summary of each Feature or, whenever needed, of each significant Work Item, introduced in Release 14."

    # Pattern to match section titles like "1 Introduction", "1.1 Overview", etc.
    heading_pattern = re.compile(r'^\d+(\.\d+)*\s+.+')

    for paragraph in document.paragraphs:
        if skipping:
            if skip_text in paragraph.text:
                skipping = False
            continue

        if heading_pattern.match(paragraph.text):
            if current_heading is not None:
                navigation_dict[current_heading] = '\n'.join(current_text)
            current_text = []
            current_heading = paragraph.text
        else:
            if current_heading is not None:
                current_text.append(paragraph.text)

    if current_heading is not None:
        navigation_dict[current_heading] = '\n'.join(current_text)

    return navigation_dict

def remove_references(text):
    # Regular expression pattern to detect the References section with possible extra lines
    reference_section_pattern = re.compile(
        r'(?:\n\s*(\[\d+\].*|TS\s*\d+\.\d+.*))+', re.MULTILINE
    )

    # Search for the reference section pattern
    match = reference_section_pattern.search(text)

    if match:
        # If a match is found, truncate the text from the match start position
        text = text[:match.start()].strip()

    return text

def remove_urls(text):
    # Regular expression pattern to detect URLs
    url_pattern = re.compile(r'https?://\S+|www\.\S+')

    # Substitute URLs with an empty string
    cleaned_text = re.sub(url_pattern, '', text)
    return cleaned_text

def remove_table_lines(text):
    # Regular expression pattern to detect lines starting with "Table" followed by various formats
    table_line_pattern = re.compile(
        r'^Table\s*[A-Za-z0-9\.\-_]+:\s*.*$', re.MULTILINE
    )

    # Remove table lines
    text = table_line_pattern.sub('', text)

    # Clean up any extra newlines left behind
    text = re.sub(r'\n\s*\n', '\n', text).strip()

    return text

def remove_non_standard_characters(text):
    # Define a regular expression pattern to match non-standard characters
    pattern = re.compile(r'[^\x20-\x7E]')  # This matches any character that is not in the standard ASCII range (0x20-0x7E)
    # Replace non-standard characters with an empty string
    cleaned_text = pattern.sub('', text)
    return cleaned_text

def process_navigation_dict(navigation_dict):
    skip_sections = ["Introduction",'Foreword', 'References',"Notifications", "Conclusion"]
    processed_dict = {}

    text_to_skip = ['void.', 'void', 'none.', 'none', "(void)"]
    
    for heading, text in navigation_dict.items():
        if text.strip():
            skip = False
            for section in skip_sections:
                if re.search(r'\b' + re.escape(section) + r'\b', heading[1]):
                    skip = True
                    break
            # Check for exact match with text_to_skip
            if not skip:
                text_stripped_lower = text.strip().lower()
                if not any(text_stripped_lower == skip_word for skip_word in text_to_skip):
                    try:
                        # Remove references from the text
                        cleaned_text = remove_references(text.strip())
                        # Remove table lines from the text
                        cleaned_text = remove_table_lines(cleaned_text)
                        # Remove URLs from the text
                        cleaned_text = remove_urls(cleaned_text)
                        cleaned_text = remove_non_standard_characters(cleaned_text)
                        
                        processed_dict[heading] = cleaned_text
                    except Exception as e:
                        print(f"Error processing heading: {heading}, error: {e}")
                        continue
    return processed_dict



def get_header_chunks(document_path):
    navigation_dict = None
    if 'rel_14.docx' in document_path:
        navigation_dict = rel14_navigation_dict(document_path)
    else:
        navigation_dict = generate_navigation_dict(document_path)
    processed_dict = process_navigation_dict(navigation_dict)

    return processed_dict



if __name__ == '__main__':
    folder_path = 'data/rel18'
    zero_item_files = []
    all_processed_dicts = {}
    
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            document_path = os.path.join(folder_path, filename)
            nav = generate_navigation_dict(document_path)
            if not nav:
                print(f"Warning: No headings detected in document '{document_path}'. Trying alternative headings detection.")
                nav = rel14_navigation_dict(document_path)

            processed_dict = process_navigation_dict(nav)
            
            if not processed_dict:
                
                zero_item_files.append(filename)
                print(f"Document '{filename}' has zero items in the processed dictionary.")
            else:
                # Convert tuple keys to strings
                all_processed_dicts[filename] = {str(k): v for k, v in processed_dict.items()}

    
    print("Documents with zero items in the processed dictionary:")
    for file in zero_item_files:
        print(file)
    
    # Save processed dictionaries to a JSON file
    with open('processed_dicts.json', 'w') as json_file:
        json.dump(all_processed_dicts, json_file, indent=4)
